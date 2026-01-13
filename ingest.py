#!/usr/bin/env python3
"""
Document ingestion script for Bahá'í Writings semantic search.
Processes .docx files, chunks by paragraphs, and stores embeddings in Pinecone.
Also supports exporting paragraphs to JSON for use with Swift embedding generator.
"""

import os
import json
import logging
import shutil
import argparse
from pathlib import Path
from typing import List, Dict, Any
from pinecone import Pinecone, ServerlessSpec
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import streamlit as st 

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class BahaiWritingsIngestor:
    def __init__(self, openai_api_key: str, pinecone_api_key: str):
        """Initialize the ingestion pipeline."""
        self.openai_client = OpenAI(api_key=openai_api_key)
        self.pc = Pinecone(api_key=pinecone_api_key)
        self.index_name = "bahai-writings"
        self.dimension = 3072  # Dimension for text-embedding-3-large
        
        # Create or get index
        try:
            # Check if index exists
            if self.index_name not in [index.name for index in self.pc.list_indexes()]:
                logger.info(f"Creating new index: {self.index_name}")
                self.pc.create_index(
                    name=self.index_name,
                    dimension=self.dimension,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
            else:
                logger.info(f"Using existing index: {self.index_name}")
            
            self.index = self.pc.Index(self.index_name)
        except Exception as e:
            logger.error(f"Error initializing Pinecone index: {e}")
            raise

    def get_author_from_filename(self, filename: str) -> str:
        """Determine the author based on filename."""
        filename_lower = filename.lower().replace('.docx', '')
        
        # Bahá'u'lláh
        bahaullah_works = [
            'kitab-i-iqan', 'hidden-words', 'gleanings-writings-bahaullah', 
            'kitab-i-aqdas-2', 'epistle-son-wolf', 'gems-divine-mysteries',
            'summons-lord-hosts', 'tablets-bahaullah', 'tabernacle-unity',
            'prayers-meditations'
        ]
        
        # 'Abdu'l-Bahá  
        abdul_baha_works = [
            'some-answered-questions', 'paris-talks', 'promulgation-universal-peace',
            'memorials-faithful', 'selections-writings-abdul-baha', 
            'secret-divine-civilization', 'travelers-narrative', 
            'will-testament-abdul-baha', 'tablets-divine-plan', 'tablet-auguste-forel'
        ]
        
        # The Báb
        bab_works = ['selections-writings-bab']
        
        # Shoghi Effendi
        shoghi_works = [
            'advent-divine-justice', 'god-passes-by', 'promised-day-come',
            'world-order-bahaullah'
        ]
        
        # Universal House of Justice (dated documents and institutional)
        uhj_works = [
            'the-institution-of-the-counsellors', 'turning-point', 'muhj-1963-1986'
        ]
        
        # Compilations
        compilation_works = ['days-remembrance', 'light-of-the-world']
        
        # Check filename against each category
        if filename_lower in bahaullah_works:
            return "Bahá'u'lláh"
        elif filename_lower in abdul_baha_works:
            return "'Abdu'l-Bahá"
        elif filename_lower in bab_works:
            return "The Báb"
        elif filename_lower in shoghi_works:
            return "Shoghi Effendi"
        elif filename_lower in uhj_works:
            return "Universal House of Justice"
        elif filename_lower in compilation_works:
            return "Compilations"
        # Handle dated documents (likely UHJ messages)
        elif filename_lower.startswith(('19', '20')) and '_' in filename_lower:
            return "Universal House of Justice"
        else:
            return "Other"

    def extract_paragraphs_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract paragraphs from a .docx file, combining short paragraphs."""
        logger.info(f"Processing document: {file_path}")
        
        doc = Document(file_path)
        raw_paragraphs = []
        
        # First, extract all non-empty paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Only keep non-empty paragraphs
                raw_paragraphs.append({
                    "text": text,
                    "original_id": i
                })
        
        # Now combine paragraphs with <min_words words
        min_words = 100
        combined_paragraphs = []
        i = 0
        
        while i < len(raw_paragraphs):
            current_text = raw_paragraphs[i]["text"]
            current_word_count = len(current_text.split())
            start_id = raw_paragraphs[i]["original_id"]
            combined_ids = [start_id]
            
            # If current paragraph has <min_words words, combine with next paragraphs
            if current_word_count < min_words and i < len(raw_paragraphs) - 1:
                j = i + 1
                while j < len(raw_paragraphs) and len(current_text.split()) < min_words:
                    next_text = raw_paragraphs[j]["text"]
                    current_text += " " + next_text
                    combined_ids.append(raw_paragraphs[j]["original_id"])
                    j += 1
                i = j  # Skip the paragraphs we just combined
            else:
                i += 1  # Move to next paragraph
            
            # Create the combined paragraph entry
            if len(combined_ids) == 1:
                document_id = f"{Path(file_path).stem}_para_{start_id}"
            else:
                document_id = f"{Path(file_path).stem}_para_{combined_ids[0]}-{combined_ids[-1]}"
            
            combined_paragraphs.append({
                "text": current_text,
                "paragraph_id": start_id,  # Use the ID of the first paragraph
                "source_file": Path(file_path).name,
                "document_id": document_id,
                "author": self.get_author_from_filename(Path(file_path).name)
            })
        
        logger.info(f"Extracted {len(raw_paragraphs)} raw paragraphs, combined into {len(combined_paragraphs)} final paragraphs")
        return combined_paragraphs

    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using OpenAI's text-embedding-3-large model."""
        try:
            response = self.openai_client.embeddings.create(
                model="text-embedding-3-large",
                input=text,
                encoding_format="float"
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            raise

    def ingest_paragraphs(self, paragraphs: List[Dict[str, Any]]):
        """Generate embeddings and store paragraphs in Pinecone."""
        logger.info(f"Generating embeddings for {len(paragraphs)} paragraphs...")
        
        vectors = []
        
        for paragraph in paragraphs:
            try:
                # Generate embedding
                embedding = self.generate_embedding(paragraph["text"])
                
                # Create vector for Pinecone
                vector = {
                    "id": paragraph["document_id"],
                    "values": embedding,
                    "metadata": {
                        "text": paragraph["text"],
                        "source_file": paragraph["source_file"],
                        "paragraph_id": paragraph["paragraph_id"],
                        "author": paragraph["author"]
                    }
                }
                vectors.append(vector)
                
                logger.info(f"Processed paragraph {paragraph['paragraph_id']}")
                
            except Exception as e:
                logger.error(f"Error processing paragraph {paragraph['paragraph_id']}: {e}")
                continue
        
        # Upsert to Pinecone index
        if vectors:
            # Pinecone has a limit on batch size, so we'll process in chunks
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i + batch_size]
                self.index.upsert(vectors=batch)
                logger.info(f"Upserted batch {i//batch_size + 1} of {(len(vectors) + batch_size - 1)//batch_size}")
            
            logger.info(f"Successfully ingested {len(vectors)} paragraphs into Pinecone")

    def move_document_to_indexed(self, file_path: str):
        """Move document from corpus to corpus_indexed folder after successful ingestion."""
        source_path = Path(file_path)
        indexed_dir = Path("./corpus_indexed")
        
        # Create corpus_indexed directory if it doesn't exist
        indexed_dir.mkdir(exist_ok=True)
        
        # Define destination path
        destination_path = indexed_dir / source_path.name
        
        try:
            shutil.move(str(source_path), str(destination_path))
            logger.info(f"Moved {source_path.name} to corpus_indexed/")
        except Exception as e:
            logger.error(f"Failed to move {source_path.name}: {e}")
            raise

    def ingest_document(self, file_path: str):
        """Complete ingestion pipeline for a single document."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Extract paragraphs
        paragraphs = self.extract_paragraphs_from_docx(file_path)
        
        # Ingest into ChromaDB
        self.ingest_paragraphs(paragraphs)
        
        # Move document to indexed folder after successful ingestion
        self.move_document_to_indexed(file_path)
        
        logger.info(f"Completed ingestion of {file_path}")

    def clear_index(self):
        """Clear all vectors from the existing index."""
        try:
            logger.info(f"Clearing all vectors from index '{self.index_name}'...")
            self.index.delete(delete_all=True)
            logger.info("Index cleared successfully")
        except Exception as e:
            logger.error(f"Error clearing index: {e}")
            raise

    def get_index_stats(self):
        """Get statistics about the index."""
        stats = self.index.describe_index_stats()
        count = stats['total_vector_count']
        logger.info(f"Index '{self.index_name}' contains {count} vectors")
        return count

def main():
    """Main ingestion function."""
    # Check for required environment variables
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    pinecone_api_key = st.secrets["PINECONE_API_KEY"]
    
    if not openai_api_key:
        raise ValueError("OPENAI_API_KEY is required in streamlit secrets")
    if not pinecone_api_key:
        raise ValueError("PINECONE_API_KEY is required in streamlit secrets")
    
    # Initialize ingestor
    ingestor = BahaiWritingsIngestor(openai_api_key, pinecone_api_key)
    
    # Check if index has existing data
    stats = ingestor.get_index_stats()
    if stats > 0:
        logger.info(f"Found {stats} existing vectors in the database.")
        
        # Ask for user confirmation before clearing
        confirmation = input(f"This will delete all {stats} existing vectors and re-ingest with new metadata. Continue? (y/N): ").strip().lower()
        
        if confirmation in ['y', 'yes']:
            logger.info("User confirmed. Clearing existing vectors from the database...")
            ingestor.clear_index()
        else:
            logger.info("Operation cancelled by user.")
            return
    else:
        logger.info("No existing vectors found. Proceeding with fresh ingestion.")
    
    # Find and process .docx files in current directory
    docx_files = list(Path("./corpus").glob("*.docx"))
    
    if not docx_files:
        logger.warning("No .docx files found in current directory")
        return
    
    for docx_file in docx_files:
        try:
            ingestor.ingest_document(str(docx_file))
        except Exception as e:
            logger.error(f"Failed to ingest {docx_file}: {e}")
    
    # Print final statistics
    ingestor.get_index_stats()

def export_paragraphs_to_json(input_dir: str = "./corpus_indexed", output_file: str = "paragraphs.json"):
    """
    Export all paragraphs from DOCX files to JSON for use with Swift embedding generator.
    This does not require any API keys.
    """
    input_path = Path(input_dir)
    docx_files = list(input_path.glob("*.docx"))

    if not docx_files:
        logger.warning(f"No .docx files found in {input_dir}")
        return

    logger.info(f"Found {len(docx_files)} documents to process")

    all_paragraphs = []

    # Create a minimal ingestor just for paragraph extraction (no API keys needed)
    class ParagraphExtractor:
        def get_author_from_filename(self, filename: str) -> str:
            """Determine the author based on filename."""
            filename_lower = filename.lower().replace('.docx', '')

            bahaullah_works = [
                'kitab-i-iqan', 'hidden-words', 'gleanings-writings-bahaullah',
                'kitab-i-aqdas-2', 'epistle-son-wolf', 'gems-divine-mysteries',
                'summons-lord-hosts', 'tablets-bahaullah', 'tabernacle-unity',
                'prayers-meditations'
            ]
            abdul_baha_works = [
                'some-answered-questions', 'paris-talks', 'promulgation-universal-peace',
                'memorials-faithful', 'selections-writings-abdul-baha',
                'secret-divine-civilization', 'travelers-narrative',
                'will-testament-abdul-baha', 'tablets-divine-plan', 'tablet-auguste-forel'
            ]
            bab_works = ['selections-writings-bab']
            shoghi_works = [
                'advent-divine-justice', 'god-passes-by', 'promised-day-come',
                'world-order-bahaullah'
            ]
            uhj_works = [
                'the-institution-of-the-counsellors', 'turning-point', 'muhj-1963-1986'
            ]
            compilation_works = ['days-remembrance', 'light-of-the-world']

            if filename_lower in bahaullah_works:
                return "Bahá'u'lláh"
            elif filename_lower in abdul_baha_works:
                return "'Abdu'l-Bahá"
            elif filename_lower in bab_works:
                return "The Báb"
            elif filename_lower in shoghi_works:
                return "Shoghi Effendi"
            elif filename_lower in uhj_works:
                return "Universal House of Justice"
            elif filename_lower in compilation_works:
                return "Compilations"
            elif filename_lower.startswith(('19', '20')) and '_' in filename_lower:
                return "Universal House of Justice"
            else:
                return "Other"

        def extract_paragraphs_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
            """Extract paragraphs from a .docx file, combining short paragraphs."""
            logger.info(f"Processing document: {file_path}")

            doc = Document(file_path)
            raw_paragraphs = []

            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    raw_paragraphs.append({"text": text, "original_id": i})

            min_words = 100
            combined_paragraphs = []
            i = 0

            while i < len(raw_paragraphs):
                current_text = raw_paragraphs[i]["text"]
                current_word_count = len(current_text.split())
                start_id = raw_paragraphs[i]["original_id"]
                combined_ids = [start_id]

                if current_word_count < min_words and i < len(raw_paragraphs) - 1:
                    j = i + 1
                    while j < len(raw_paragraphs) and len(current_text.split()) < min_words:
                        next_text = raw_paragraphs[j]["text"]
                        current_text += " " + next_text
                        combined_ids.append(raw_paragraphs[j]["original_id"])
                        j += 1
                    i = j
                else:
                    i += 1

                if len(combined_ids) == 1:
                    document_id = f"{Path(file_path).stem}_para_{start_id}"
                else:
                    document_id = f"{Path(file_path).stem}_para_{combined_ids[0]}-{combined_ids[-1]}"

                combined_paragraphs.append({
                    "text": current_text,
                    "paragraph_id": start_id,
                    "source_file": Path(file_path).name,
                    "document_id": document_id,
                    "author": self.get_author_from_filename(Path(file_path).name)
                })

            logger.info(f"Extracted {len(raw_paragraphs)} raw paragraphs, combined into {len(combined_paragraphs)} final paragraphs")
            return combined_paragraphs

    extractor = ParagraphExtractor()

    for docx_file in sorted(docx_files):
        try:
            paragraphs = extractor.extract_paragraphs_from_docx(str(docx_file))
            all_paragraphs.extend(paragraphs)
            logger.info(f"Processed {docx_file.name}: {len(paragraphs)} paragraphs")
        except Exception as e:
            logger.error(f"Failed to process {docx_file}: {e}")

    # Write to JSON
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_paragraphs, f, ensure_ascii=False, indent=2)

    logger.info(f"Exported {len(all_paragraphs)} paragraphs to {output_file}")
    return all_paragraphs


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Bahá'í Writings ingestion and export tool")
    parser.add_argument('--export-json', type=str, metavar='OUTPUT_FILE',
                        help='Export paragraphs to JSON file (no API keys required)')
    parser.add_argument('--input-dir', type=str, default='./corpus_indexed',
                        help='Input directory containing .docx files (default: ./corpus_indexed)')

    args = parser.parse_args()

    if args.export_json:
        export_paragraphs_to_json(args.input_dir, args.export_json)
    else:
        main()